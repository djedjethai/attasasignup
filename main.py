import os
import platform
import pandas as pd
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

# from docx.oxml.ns import nsdecls
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from utility import replace_title, may_empty

# Load the responses from the Excel file
input_file = "./data/responses.xlsx"
responses = pd.read_excel(input_file)

# Define the template file (can be an empty form with placeholders)
template_file = "./forms/formTest.docx"

placeholder = ''

# Process each row and generate a form
for index, row in responses.iterrows():
    # Load the template
    doc = Document(template_file)
    
    # Replace placeholders with actual data
    for paragraph in doc.paragraphs:
        paragraph.style.font.name = 'TH Sarabun New'
        paragraph.style.font.size = Pt(12)
     
        
        print('the paragraph is: ', paragraph.text)
        for col_index, (key, value) in enumerate(row.items()):  # Use enumerate
            print(f"Processing column {col_index}: key = {key}, value = {value}")

            if key == 'เพศ':    
                paragraph = replace_title(paragraph, str(value)) 
                
            elif key == 'ชื่อ':  # Correct syntax for equality check
                placeholder = '$2'
                paragraph.text = paragraph.text.replace(placeholder, str(value))

            elif key == 'นามสกุล':  # Correct syntax for equality check
                placeholder = '$3'
                paragraph.text = paragraph.text.replace(placeholder, str(value))

            elif key == 'วัน/เดือน/ปี (พ.ศ.) เกิด':  # Correct syntax for equality check
                placeholder = '$4'
                paragraph.text = paragraph.text.replace(placeholder, str(value))
            
            elif key == 'ที่อยู่ปัจจุบัน เลขที่':  # Correct syntax for equality check
                placeholder = '$5'
                paragraph.text = paragraph.text.replace(placeholder, str(value))

            elif key == 'หมู่ที่':  # Correct syntax for equality check
                placeholder = '$6'
                paragraph = may_empty(paragraph, value, placeholder)

            elif key == 'หมู่บ้าน':  # Correct syntax for equality check
                placeholder = '$7'
                paragraph = may_empty(paragraph, value, placeholder)

            elif key == 'ซอย':  # Correct syntax for equality check
                placeholder = '$8'
                paragraph = may_empty(paragraph, value, placeholder)

            elif key == 'ถนน':  # Correct syntax for equality check
                placeholder = '$9'
                paragraph = may_empty(paragraph, value, placeholder)

            elif key == 'แขวง':  # Correct syntax for equality check
                placeholder = '$10'
                paragraph.text = paragraph.text.replace(placeholder, str(value))

            elif key == '  กรุณาตรวจสอบคำตอบของคุณก่อนที่จะส่ง  ':  # Correct syntax for equality check
                img_path = './data/images.png'
                place = 'picture'

                # Split the paragraph into runs and look for the placeholder
                for run in paragraph.runs:
                    print('Found placeholder run.text is :', run.text)
                    run.text = run.text.strip()
                    if place in run.text:  # Find the run containing the placeholder
                        # Split the text at the placeholder
                        text_before = run.text.split(place)[0]
                        text_after = run.text.split(place)[1]
                        
                        # Clear the existing run text and add new text before and after the placeholder
                        run.text = text_before
                        new_run = paragraph.add_run()  # Add a new run for the image
                        new_run.add_picture(img_path, width=Inches(1), height=Inches(1))
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

                        # Add the remaining text after the placeholder
                        run = paragraph.add_run()
                        run.text = text_after
                        
                        # picture_added = True  # Mark the picture as added
                        break  # Exit the loop after adding the picture
                
            else:
                print(f"In else {col_index}: key = {key}, value = {value}")

        # # NOTE 1 ok
        # for run in paragraph.runs:
        #     # run.font.name = 'Liberation Sans'
        #     run.font.name = 'TH Sarabun New'
        #     run.font.size = Pt(17)
       
           
    # Save the filled form
    output_directory = "./results/"
    output_file = f"{output_directory}form_{index + 1}.docx"
    # output_file = f"form_{index + 1}.docx"  # Name the file based on the row index
    doc.save(output_file)

print("Forms generated successfully!")


